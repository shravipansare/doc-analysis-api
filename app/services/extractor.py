import os
import logging
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ExtractResult:
    text: str
    page_count: Optional[int]
    ocr_used: bool
    word_count: int = field(init=False)

    def __post_init__(self):
        self.word_count = len(self.text.split()) if self.text else 0


def extract_pdf(file_path: str) -> ExtractResult:
    """Extract text from PDF using PyMuPDF. Falls back to OCR if needed."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        page_count = len(doc)
        text_parts = []

        for page in doc:
            # Multi-column sort
            text_parts.append(page.get_text("text", sort=True))
            
            # Table extraction to Markdown
            tabs = page.find_tables()
            if tabs and tabs.tables:
                for table in tabs.tables:
                    md_table = table.to_markdown()
                    if md_table:
                        text_parts.append(f"\n[TABLE EXTRACTION]\n{md_table}\n")

        doc.close()
        full_text = "\n\n".join(text_parts).strip()

        # If minimal text extracted, likely a scanned PDF — use OCR
        if len(full_text) < 100:
            logger.info("PDF has minimal native text, switching to OCR...")
            return _ocr_pdf(file_path, page_count)

        return ExtractResult(text=full_text, page_count=page_count, ocr_used=False)

    except Exception as e:
        logger.error(f"PyMuPDF extraction failed: {e}")
        raise


def _ocr_pdf(file_path: str, page_count: int) -> ExtractResult:
    """Convert PDF pages to images and apply Tesseract OCR."""
    try:
        from pdf2image import convert_from_path
        import pytesseract
        from PIL import Image

        images = convert_from_path(file_path, dpi=300)
        text_parts = []

        for img in images:
            # Preprocess: convert to grayscale for better OCR accuracy
            gray = img.convert("L")
            text = pytesseract.image_to_string(gray, config="--psm 1")
            text_parts.append(text)

        full_text = "\n\n".join(text_parts).strip()
        return ExtractResult(text=full_text, page_count=page_count, ocr_used=True)

    except Exception as e:
        logger.error(f"OCR for PDF failed: {e}")
        raise


def extract_docx(file_path: str) -> ExtractResult:
    """Extract text from DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extract regular paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract text from tables
        for table in doc.tables:
            if not table.rows:
                continue
            text_parts.append("\n[TABLE EXTRACTION]")
            for i, row in enumerate(table.rows):
                # Clean cell text (remove newlines to prevent markdown breakage)
                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                row_text = "| " + " | ".join(row_cells) + " |"
                text_parts.append(row_text)
                if i == 0:
                    divider = "| " + " | ".join(["---"] * len(row_cells)) + " |"
                    text_parts.append(divider)
            text_parts.append("\n")

        full_text = "\n\n".join(text_parts).strip()
        return ExtractResult(text=full_text, page_count=None, ocr_used=False)

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise


def extract_image(file_path: str) -> ExtractResult:
    """Extract text from image using Pillow preprocessing + Tesseract OCR."""
    try:
        import pytesseract
        from PIL import Image, ImageFilter, ImageEnhance

        img = Image.open(file_path)

        # Convert to RGB if needed (e.g., TIFF might be CMYK)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        # Preprocessing pipeline for better OCR accuracy
        # 1. Convert to grayscale
        gray = img.convert("L")

        # 2. Enhance contrast
        enhancer = ImageEnhance.Contrast(gray)
        enhanced = enhancer.enhance(2.0)

        # 3. Apply slight sharpening
        sharpened = enhanced.filter(ImageFilter.SHARPEN)

        # Run Tesseract with page segmentation mode 1 (auto page seg + OSD)
        text = pytesseract.image_to_string(sharpened, config="--psm 1 --oem 3")
        text = text.strip()

        return ExtractResult(text=text, page_count=1, ocr_used=True)

    except Exception as e:
        logger.error(f"Image OCR extraction failed: {e}")
        raise


def extract_text(file_path: str, file_type: str) -> ExtractResult:
    """
    Main dispatcher: routes file to appropriate extractor.

    Args:
        file_path: Absolute path to the temporary file
        file_type: 'pdf', 'docx', or 'image'

    Returns:
        ExtractResult with text, page_count, ocr_used
    """
    if file_type == "pdf":
        return extract_pdf(file_path)
    elif file_type == "docx":
        return extract_docx(file_path)
    elif file_type == "image":
        return extract_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
